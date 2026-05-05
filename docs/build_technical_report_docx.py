#!/usr/bin/env python3
"""
Build the CS581 Signature Project technical report as a .docx file.

Usage (from repo root):
    python3 -m venv /tmp/docx-venv
    /tmp/docx-venv/bin/pip install python-docx
    /tmp/docx-venv/bin/python docs/build_technical_report_docx.py

Output: docs/technical-report.docx
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches
from pathlib import Path


# ---------- helpers ----------

def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F) if level <= 2 else RGBColor(0x2E, 0x4E, 0x7E)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        render_inline(p, item)


def add_numbered(doc: Document, items: list) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        render_inline(p, item)


def render_inline(paragraph, text: str) -> None:
    """Render a string with **bold**, *italic*, and `code` spans."""
    import re

    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_table(doc: Document, headers: list, rows: list) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(hdr[i], "1F3A5F")
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            render_inline(p, val)
            for r in p.runs:
                r.font.size = Pt(10)
    doc.add_paragraph()


def add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------- build the document ----------

def build(out_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Title page ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Secure Cloud-Native Application\nDeployment on AWS EKS")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Technical Report")
    r.italic = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line, size in [
        ("CS581 — Cloud Security", 13),
        ("Signature Project", 13),
        ("San Francisco Bay University", 12),
        ("Spring 2026", 12),
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(size)

    doc.add_paragraph()
    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = scope.add_run(
        "Design, Implementation, and Security Hardening of a "
        "Cloud-Native Fintech Application on Amazon EKS"
    )
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_page_break()

    # ---- Executive Summary ----
    add_heading(doc, "Executive Summary", 1)
    add_para(doc,
        "This project delivers a production-grade, secure deployment of a multi-tier "
        "fintech web application on Amazon EKS. The architecture applies defense-in-depth "
        "across seven security domains — network, identity, data, container, runtime, "
        "monitoring, and incident response — and is fully reproducible via Infrastructure as Code."
    )
    add_para(doc, "The implementation includes:")
    add_bullets(doc, [
        "A hardened AWS network (VPC with public/private subnet separation, NAT egress, "
        "restrictive Security Groups and NACLs, and VPC Flow Logs)",
        "An EKS cluster with control-plane audit logging, KMS-encrypted secrets, and "
        "managed node groups in private subnets",
        "Least-privilege AWS IAM roles for cluster infrastructure, worker nodes, pod "
        "workloads (IRSA), and three human-user personas (admin, developer, viewer)",
        "Kubernetes RBAC with namespace isolation and Pod Security Standards (restricted profile)",
        "Encrypted data at rest (KMS on EBS, EKS secrets, Secrets Manager, ECR) and in "
        "transit (TLS 1.3 via ACM on the ALB)",
        "ECR with immutable tags, scan-on-push, and Amazon Inspector enhanced scanning",
        "Trivy CI/CD scanning of source, container images, and Kubernetes manifests",
        "GuardDuty with EKS audit-log monitoring, CloudWatch Container Insights, and "
        "alarms routed to SNS",
        "Three validated threat-simulation scenarios demonstrating the controls work end to end",
    ])
    add_para(doc,
        "All infrastructure is provisioned by Terraform and all workloads by Kubernetes "
        "manifests, with GitHub Actions orchestrating image builds and vulnerability scanning. "
        "The entire system can be deployed or destroyed in approximately 20 minutes."
    )
    add_divider(doc)

    # ---- 1. Architecture ----
    add_heading(doc, "1. Architecture Design and Justification", 1)

    add_heading(doc, "1.1 System Context", 2)
    add_para(doc, "A fintech startup requires a cloud-native web application that:")
    add_bullets(doc, [
        "Handles sensitive user and transaction data",
        "Deploys to AWS EKS",
        "Resists common cloud and container threats",
        "Complies with foundational security controls: least privilege, encryption, "
        "auditability, and incident response",
    ])

    add_heading(doc, "1.2 Architecture Overview", 2)
    add_para(doc, "The solution is organized in three conceptual planes:")
    add_numbered(doc, [
        "**Network plane** — A single-region VPC (`10.0.0.0/16`) spread across three "
        "Availability Zones with paired public/private subnets. Public subnets host the "
        "ALB and NAT gateways; private subnets host EKS worker nodes and the EKS-managed "
        "control plane ENIs.",
        "**Compute plane** — EKS managed node group (2–4 × t3.medium) running in private "
        "subnets. All pods are scheduled here; no pod is ever scheduled in a public subnet.",
        "**Security and observability plane** — IAM (roles and IRSA), KMS (customer-managed "
        "keys), Secrets Manager, CloudWatch Logs, GuardDuty, and SNS are layered on top.",
    ])
    note = doc.add_paragraph()
    r = note.add_run(
        "Figure: Architecture diagram (see docs/architecture-diagram.py → architecture.png). "
        "Generate the PNG and paste it below before submitting."
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    add_heading(doc, "1.3 Design Decisions and Security Justification", 2)
    add_table(doc,
        ["Decision", "Alternative Considered", "Security Rationale"],
        [
            ["EKS nodes in private subnets only", "Nodes in public subnets",
             "Nodes have no public IPs; inbound traffic terminates at the ALB. Eliminates direct node exposure."],
            ["Multi-AZ VPC (3 AZs)", "Single AZ",
             "High availability; AZ outage does not take down the app."],
            ["NAT Gateway (public subnets)", "Public nodes for egress",
             "Nodes reach AWS APIs (ECR, STS, Secrets Manager) without being reachable from the Internet."],
            ["Customer-managed KMS key", "AWS-managed keys",
             "Full control over key policy, rotation, and access audit."],
            ["Managed node group with Launch Template", "Self-managed EC2",
             "EKS owns lifecycle (upgrades, AMI patching). Launch Template enforces encrypted EBS."],
            ["IRSA for pod workloads", "Instance profile inheritance",
             "Per-pod AWS permissions via OIDC — no shared credentials, no static keys."],
            ["Three RBAC personas (admin/dev/viewer)", "Single developer role",
             "Matches real-world separation of duties. Audit trail includes the specific role used."],
            ["Pod Security Standard `restricted`", "`baseline` or no PSS",
             "Blocks privileged pods, host-namespace use, and capability additions at admission."],
            ["ALB Ingress with ACM TLS 1.3", "Self-signed or cert-manager",
             "AWS-managed certs, automatic rotation, enforced modern cipher policy."],
            ["Immutable ECR image tags", "Mutable tags",
             "Deployed image contents cannot change under the running pod."],
            ["GuardDuty with EKS audit-log monitoring", "VPC Flow Logs review only",
             "Automated, continuous analysis versus manual forensics."],
            ["Terraform for all infrastructure", "Click-ops / eksctl",
             "Reviewable, reproducible, diffable changes. Drift detection via `terraform plan`."],
        ],
    )

    add_heading(doc, "1.4 Threat Model Summary (STRIDE)", 2)
    add_table(doc,
        ["Threat", "Affected Component", "Control"],
        [
            ["**S**poofing user identity", "API access",
             "IAM role authentication + Kubernetes RBAC groups"],
            ["**T**ampering with data in transit", "Client ↔ ALB ↔ pod",
             "TLS 1.3 at ALB; in-cluster TLS for kubelet/API"],
            ["**T**ampering with data at rest", "EBS, Secrets Manager, ECR",
             "KMS encryption on all three"],
            ["**R**epudiation", "Any admin action",
             "CloudTrail (AWS) + EKS audit log (Kubernetes)"],
            ["**I**nformation disclosure", "Secret material",
             "Secrets Manager with IRSA; no secrets in images or manifests"],
            ["**D**enial of service", "ALB / nodes",
             "ALB WAF-ready; HPA-ready node group; resource quotas on pods"],
            ["**E**levation of privilege", "Pod → host, user → admin",
             "PSS `restricted`, RBAC least privilege, `readOnlyRootFilesystem`"],
        ],
    )
    add_divider(doc)

    # ---- 2. Implementation ----
    add_heading(doc, "2. Implementation Summary", 1)

    add_heading(doc, "2.1 Infrastructure (Terraform)", 2)
    add_para(doc, "Six Terraform modules, each with a single responsibility:")
    add_table(doc,
        ["Module", "Resources"],
        [
            ["`vpc`", "VPC, 6 subnets, IGW, 3 NAT GWs, route tables, VPC Flow Logs + CloudWatch log group"],
            ["`iam`", "Cluster role, node role, OIDC provider prerequisites, three RBAC-persona roles, `eks:DescribeCluster` policy"],
            ["`security`", "EKS cluster SG, node SG, NACLs, customer-managed KMS key + alias"],
            ["`eks`", "EKS cluster (with encryption config), managed node group, launch template (encrypted gp3 EBS), OIDC provider, CloudWatch log group for control-plane logs"],
            ["`container-security`", "ECR repos (frontend, backend) with `IMMUTABLE` tags, `scan_on_push`, KMS encryption, Enhanced Scanning registry config, lifecycle policies"],
            ["`monitoring`", "GuardDuty detector with EKS audit-log + EBS malware scanning, SNS alarm topic (KMS-encrypted), CloudWatch alarms (failed auth, node CPU, new GuardDuty findings), Container Insights add-on"],
        ],
    )
    add_para(doc, "The root module also provisions:")
    add_bullets(doc, [
        "`aws_secretsmanager_secret` for backend app secrets (KMS-encrypted, 7-day recovery window)",
        "IRSA role for the backend ServiceAccount (`system:serviceaccount:fintech-app:app-backend`) "
        "with a narrowly scoped policy (Secrets Manager Get/Describe on one ARN, KMS Decrypt on one key)",
    ])

    add_heading(doc, "2.2 Kubernetes Workloads", 2)
    add_table(doc,
        ["Resource Type", "Manifest", "Purpose"],
        [
            ["Namespace", "`namespaces/namespaces.yaml`", "`fintech-app` (PSS `restricted`), `monitoring` (PSS `baseline`)"],
            ["ClusterRole", "`rbac/cluster-roles.yaml`", "`cluster-admin-role`, `developer-role`, `viewer-role`"],
            ["Role", "`rbac/app-roles.yaml`", "`app-deployer`, `app-viewer`, `monitoring-operator`"],
            ["Bindings", "`rbac/role-bindings.yaml`", "Group → role (cluster-wide + namespaced)"],
            ["ConfigMap", "`rbac/aws-auth-configmap.yaml`", "IAM role ARN → K8s group"],
            ["ServiceAccount", "`rbac/service-accounts.yaml`", "`app-backend` (IRSA) and `app-frontend` (no IRSA)"],
            ["Deployment", "`deployments/backend.yaml`", "2× Flask backend, non-root UID 1000, read-only root FS"],
            ["Deployment", "`deployments/frontend.yaml`", "2× nginx-unprivileged, non-root UID 101, listens on 8080"],
            ["StatefulSet", "`deployments/postgres.yaml`", "Single replica, non-root UID 999, encrypted gp3 volume"],
            ["Services", "`services/services.yaml`", "All ClusterIP (internal); database is a headless service"],
            ["Ingress", "`ingress/app-ingress-acm.yaml`", "ALB with ACM TLS 1.3, SSL redirect, routes `/api` → backend and `/` → frontend"],
            ["NetworkPolicy", "`network-policies/deny-all.yaml`", "Default deny ingress + egress, allow DNS egress"],
            ["NetworkPolicy", "`network-policies/frontend-policy.yaml`", "Frontend ← ALB, Frontend → backend (5000)"],
            ["NetworkPolicy", "`network-policies/backend-policy.yaml`", "Backend ← frontend; backend → DB (5432) + AWS APIs (443); DB ← backend only"],
        ],
    )

    add_heading(doc, "2.3 CI/CD", 2)
    add_bullets(doc, [
        "`trivy-scan.yml` — filesystem, container-image, and Kubernetes-config scans on every pull request; results uploaded as SARIF to GitHub Code Scanning.",
        "`build-and-push.yml` — on push to `main` or version tag, authenticates to AWS via OIDC (no long-lived keys), builds images, enforces a final CRITICAL-severity Trivy gate, and pushes to ECR.",
    ])

    add_heading(doc, "2.4 Application", 2)
    add_bullets(doc, [
        "**Backend** — Flask + SQLAlchemy + Flask-JWT-Extended, reads config from environment or AWS Secrets Manager via IRSA",
        "**Frontend** — React 18 + Vite SPA with protected routes, served by nginx-unprivileged",
        "**Database** — PostgreSQL 16-alpine running as a StatefulSet with encrypted gp3 storage",
    ])
    add_divider(doc)

    # ---- 3. Security Controls ----
    add_heading(doc, "3. Security Controls Implemented", 1)

    add_heading(doc, "3.1 Network Security", 2)
    add_bullets(doc, [
        "**VPC segmentation**: public subnets (ALB, NAT) vs private subnets (nodes, control plane). Nodes have no public IPs.",
        "**Security Groups**: cluster SG allows only control-plane traffic; node SG allows only cluster ↔ node and ALB ↔ node.",
        "**NACLs**: stateless defense-in-depth at the subnet boundary.",
        "**VPC Flow Logs**: every packet's metadata sent to CloudWatch, encrypted.",
        "**Network Policies**: default deny-all ingress and egress in `fintech-app`, with explicit allow rules for frontend → backend → DB and backend → AWS APIs. DNS egress allowed namespace-wide to `kube-system/kube-dns`.",
    ])

    add_heading(doc, "3.2 Identity and Access Management", 2)
    add_bullets(doc, [
        "**AWS IAM least privilege** — Cluster role: `AmazonEKSClusterPolicy` + `AmazonEKSVPCResourceController`. Node role: worker + CNI + ECR read-only. Persona roles: `eks:DescribeCluster` only.",
        "**IRSA** — The backend ServiceAccount assumes a role whose trust policy is scoped to exactly one OIDC subject (`system:serviceaccount:fintech-app:app-backend`). The role policy allows only `secretsmanager:Get/DescribeSecret` on one ARN and `kms:Decrypt` on one key.",
        "**Kubernetes RBAC** — Three ClusterRoles plus three namespaced Roles; destructive actions (delete, exec on RBAC objects) are forbidden to developers and viewers.",
        "**`aws-auth` ConfigMap** — Explicit IAM ARN → K8s group mapping. Node-group mapping preserved; human users have no implicit mapping.",
    ])

    add_heading(doc, "3.3 Data Security", 2)
    add_bullets(doc, [
        "**At rest** — KMS customer-managed key encrypts EKS secrets, EBS volumes, Secrets Manager, ECR images, and the SNS topic.",
        "**In transit** — ALB serves TLS 1.3 with `ELBSecurityPolicy-TLS13-1-2-2021-06`; SSL redirect forces all HTTP to HTTPS. In-cluster traffic uses kubelet's TLS.",
        "**Secrets** — Backend pulls the `fintech-secure-dev-backend` Secrets Manager bundle via IRSA on startup. No secrets live in manifests, environment variables committed to Git, or container images.",
    ])

    add_heading(doc, "3.4 Container Security", 2)
    add_bullets(doc, [
        "**Build time** — Multi-stage Dockerfiles, minimal slim/alpine bases, dedicated non-root users.",
        "**Registry** — ECR with immutable tags, KMS encryption, scan-on-push, and Amazon Inspector enhanced continuous scanning.",
        "**CI/CD** — Trivy scans source, images, and Kubernetes manifests; pipeline fails on unfixed HIGH/CRITICAL CVEs.",
        "**Runtime** — Pod Security Standard `restricted` enforced at the namespace; every pod sets `runAsNonRoot`, `readOnlyRootFilesystem`, `allowPrivilegeEscalation: false`, drops ALL capabilities, and uses `RuntimeDefault` seccomp.",
    ])

    add_heading(doc, "3.5 Monitoring and Incident Response", 2)
    add_bullets(doc, [
        "**Logs** — EKS control plane (api, audit, authenticator, controllerManager, scheduler) + VPC Flow Logs + Container Insights, all retained 30 days in CloudWatch.",
        "**Threat detection** — GuardDuty with EKS audit-log analysis, EBS malware scanning, and DNS/VPC/CloudTrail anomaly detection.",
        "**Alerting** — Three CloudWatch alarms → SNS (KMS-encrypted) → email subscription, covering failed API authentication, node CPU exhaustion, and new GuardDuty findings.",
        "**Runbook** — Phase 9 threat-simulation document captures incident-response steps for each scenario.",
    ])
    add_divider(doc)

    # ---- 4. Threat model + mitigation ----
    add_heading(doc, "4. Threat Model and Mitigation Strategies", 1)

    add_para(doc,
        "Three attack scenarios were executed against the deployed cluster to validate the "
        "controls end to end. Full transcripts are captured in docs/phase9-threat-simulation.md "
        "and the screenshots PDF."
    )

    add_heading(doc, "4.1 Scenario 1 — Unauthorized API Access", 2)
    add_para(doc, "**Attack.** A compromised `viewer` IAM role attempts `kubectl delete deployment backend -n fintech-app`.")
    add_para(doc, "**Detection and prevention.** Blocked synchronously by Kubernetes RBAC with a 403 Forbidden response. The denial is captured in the EKS audit log with the user, verb, resource, and timestamp. If more than 10 such denials occur in 5 minutes the `failed-auth` CloudWatch alarm fires and pages the SNS topic.")
    add_para(doc, "**Mitigation.** Least-privilege ClusterRole grants only read verbs; the developer and viewer roles cannot execute destructive actions. Regular review of the `aws-auth` ConfigMap and `kubectl auth can-i --list` output ensures no role drift.")
    add_para(doc, "**Incident response.** Identify the source IAM principal from the audit log → rotate the suspected role's trust policy → review CloudTrail for lateral use → if confirmed, disable the AWS user or role and force credential rotation.")

    add_heading(doc, "4.2 Scenario 2 — Privilege Escalation via Privileged Pod", 2)
    add_para(doc, "**Attack.** A malicious manifest attempts to create a pod with `privileged: true`, `hostPID: true`, `hostNetwork: true`, and the `SYS_ADMIN` capability.")
    add_para(doc, "**Detection and prevention.** Blocked at admission by Pod Security Standards with a detailed violation message enumerating every policy the pod breaks. No pod object is created; no image is pulled; no runtime detection is required.")
    add_para(doc, "**Mitigation.** The `fintech-app` namespace has `pod-security.kubernetes.io/enforce: restricted`. Every application pod is written to pass the restricted profile without exception, so PSS can be enforced rather than audited. Escape-hatch namespaces are forbidden.")
    add_para(doc, "**Incident response.** The admission denial is itself the alert. Review the audit log for the identity that submitted the pod; correlate with CI/CD pipeline runs (did a malicious PR reach apply?); if the source is a bot/service account, rotate its credentials.")

    add_heading(doc, "4.3 Scenario 3 — Root Credential Usage", 2)
    add_para(doc, "**Attack.** An administrator issues AWS API calls using root account credentials.")
    add_para(doc, "**Detection.** GuardDuty raises a `Policy:IAMUser/RootCredentialUsage` finding within minutes. The `guardduty-findings` CloudWatch alarm fires and emails the on-call address via SNS.")
    add_para(doc, "**Mitigation.** The three persona IAM roles (platform-admin, developer, viewer) cover every legitimate operation, so root is never required. MFA on the root account, and strict key hygiene, mean a finding indicates either a test or an incident — both warrant investigation.")
    add_para(doc, "**Incident response.** Confirm the actor is an authorized admin testing; if not, immediately rotate root credentials → audit CloudTrail for the full blast radius → open an incident ticket and log the timeline.")

    add_para(doc,
        "All three scenarios demonstrate defense in depth in action: scenarios 1 and 2 were "
        "prevented by synchronous controls (RBAC, PSS admission); scenario 3 was detected and "
        "escalated asynchronously (GuardDuty + CloudWatch + SNS). Removing any single control "
        "would have let the corresponding attack succeed, which is exactly why they are layered."
    )
    add_divider(doc)

    # ---- 5. Evaluation ----
    add_heading(doc, "5. Evaluation Against Project Requirements", 1)
    add_table(doc,
        ["Requirement", "Evidence"],
        [
            ["VPC with public/private subnets", "`terraform/modules/vpc/`"],
            ["EKS with managed node groups", "`terraform/modules/eks/`"],
            ["Terraform for all infrastructure", "6 Terraform modules + root"],
            ["Multi-tier K8s app (frontend + backend + DB)", "`kubernetes/deployments/`"],
            ["IAM least privilege + IRSA", "`terraform/modules/iam/` + root `main.tf`"],
            ["Kubernetes RBAC", "`kubernetes/rbac/`"],
            ["Security Groups + NACLs + private subnets", "`terraform/modules/security/`"],
            ["Network Policies", "`kubernetes/network-policies/`"],
            ["Data encryption at rest", "KMS on EBS, EKS secrets, Secrets Manager, ECR, SNS"],
            ["TLS in transit", "ALB + ACM with TLS 1.3 policy"],
            ["Secrets Manager integration", "`aws_secretsmanager_secret` + IRSA policy"],
            ["Image scanning (ECR + Trivy)", "`terraform/modules/container-security/` + `.github/workflows/`"],
            ["Non-root containers, minimal base", "Hardened Dockerfiles + `securityContext`"],
            ["Pod Security Standards", "`fintech-app` namespace labels"],
            ["CloudWatch logs + Kubernetes audit logs", "EKS cluster logging + `/aws/eks/.../cluster` log group"],
            ["GuardDuty", "`terraform/modules/monitoring/`"],
            ["Prometheus + Grafana (optional)", "`kubernetes/monitoring/prometheus-values.yaml`"],
            ["≥ 2 threat simulations", "3 scenarios documented with transcripts"],
            ["Incident response documentation", "Each Phase 9 scenario has an IR section"],
        ],
    )
    add_divider(doc)

    # ---- 6. Lessons learned ----
    add_heading(doc, "6. Lessons Learned", 1)
    lessons = [
        ("Defense in depth is non-negotiable.",
         "No single control blocked every attack in Phase 9. RBAC blocked the unauthorized "
         "action, PSS blocked the privileged pod, and GuardDuty caught the misconfiguration. "
         "Remove any one of those layers and a realistic attacker gets further."),
        ("IRSA is the correct answer to \"how do pods get AWS access.\"",
         "Every alternative (instance profiles, static keys, shared roles) either violates "
         "least privilege or creates credential-rotation pain. The one-time cost of wiring "
         "OIDC pays back immediately."),
        ("Immutable tags + continuous scanning beats tag discipline.",
         "Early iterations mutated `:latest` and relied on developers remembering to re-scan. "
         "Immutable tags and Inspector Enhanced Scanning made \"the image running in prod\" "
         "and \"the image we scanned\" mathematically the same thing."),
        ("Pod Security Standards are a better default than PSP or custom admission webhooks.",
         "PSS ships with the cluster, has three clear tiers (privileged, baseline, restricted), "
         "and the violation messages point directly at the offending field. No webhook to deploy, "
         "no policy DSL to learn."),
        ("Terraform modules over flat config.",
         "Splitting the project into six modules (vpc, iam, security, eks, container-security, "
         "monitoring) made each concern reviewable in isolation. Adding Phase 7 and 8 was two "
         "new modules, not a refactor of existing code."),
        ("NetworkPolicies are ingress *and* egress.",
         "An early version only had ingress rules; the backend could still make outbound "
         "connections anywhere. Adding explicit egress rules (to DB, to AWS APIs on 443 only) "
         "closed a real lateral-movement path."),
        ("Alerting is only as good as the runbook behind it.",
         "CloudWatch firing on \"10 failed auths\" is noise unless someone knows what to do "
         "with it. The runbook written alongside each alarm is as important as the alarm itself."),
        ("Cost-discipline matters at the educational scale too.",
         "The full stack costs roughly $7/day. Keeping the cluster up 24/7 burns through "
         "project budget quickly; `terraform destroy` after each work session is the cheapest "
         "habit to form."),
    ]
    for i, (title, body) in enumerate(lessons, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {title} ")
        r.bold = True
        render_inline(p, body)

    add_divider(doc)

    # ---- Appendix ----
    add_heading(doc, "Appendix A — File Index", 1)
    add_bullets(doc, [
        "`terraform/` — All infrastructure modules and root configuration",
        "`kubernetes/` — All cluster-level manifests, in subdirectories matching the deployment order",
        "`backend/` — Flask application source and hardened Dockerfile",
        "`frontend/` — React application source and hardened Dockerfile",
        "`.github/workflows/` — CI/CD pipelines (Trivy scan, ECR build-and-push)",
        "`docs/` — Phase-level deep dives, architecture diagram source, and this report",
    ])

    add_heading(doc, "Appendix B — Command Reference", 1)
    add_table(doc,
        ["Task", "Command"],
        [
            ["Provision everything", "`cd terraform && terraform apply`"],
            ["Configure kubectl", "`aws eks update-kubeconfig --region us-west-2 --name fintech-secure-dev`"],
            ["Apply all K8s manifests", "`kubectl apply -R -f kubernetes/` (after substituting placeholders)"],
            ["View audit log", "`aws logs tail /aws/eks/fintech-secure-dev/cluster --since 10m`"],
            ["List GuardDuty findings", "`aws guardduty list-findings --detector-id <id>`"],
            ["Tear down everything", "`kubectl delete -R -f kubernetes/ && cd terraform && terraform destroy`"],
        ],
    )

    add_heading(doc, "Appendix C — References", 1)
    add_bullets(doc, [
        "AWS EKS Best Practices Guide — https://aws.github.io/aws-eks-best-practices/",
        "Kubernetes Pod Security Standards — https://kubernetes.io/docs/concepts/security/pod-security-standards/",
        "CIS Amazon EKS Benchmark v1.5",
        "NIST SP 800-190 — Application Container Security Guide",
        "OWASP Kubernetes Top 10",
    ])

    doc.save(out_path)


if __name__ == "__main__":
    out = Path(__file__).parent / "technical-report.docx"
    build(out)
    print(f"Wrote {out}")
